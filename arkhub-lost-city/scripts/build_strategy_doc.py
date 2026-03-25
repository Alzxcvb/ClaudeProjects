#!/usr/bin/env python3
"""
Build the Arkhub discovery strategy Word document.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).parent.parent / "docs" / "discovery_strategy.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# ── Page margins ───────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)

# ── Helpers ────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)  # dark indigo
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x28, 0x3E, 0xA7)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x37, 0x47, 0xC3)
    return p

def body(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
    p.add_run(text)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p

def table_row(tbl, cells, bold=False):
    row = tbl.add_row()
    for i, txt in enumerate(cells):
        cell = row.cells[i]
        cell.text = txt
        if bold:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    return row

def shade_row(row, hex_color="D9D9D9"):
    """Shade a table row background."""
    for cell in row.cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(shading)


# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("ARKHUB: LOST CITY DISCOVERY")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run("Archaeological Site Discovery Strategy — Peru")
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x55, 0x5F, 0x6E)

sub_p2 = doc.add_paragraph()
sub_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = sub_p2.add_run("NS Archaeology Hackathon · March 2026")
run3.font.size = Pt(11)
run3.font.color.rgb = RGBColor(0x75, 0x82, 0x8F)

doc.add_paragraph()
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════
h1("Executive Summary")
body(
    "This document presents a multi-layer analytical strategy for identifying undiscovered "
    "archaeological sites in Peru at 1 km² resolution. Using 34 verified known-site coordinates "
    "and 118 curated academic papers assembled from OpenAlex and Crossref, we apply six "
    "independent predictive layers to generate a composite score for every 1 km² cell "
    "across four priority study regions.\n\n"
    "The strategy does not require satellite imagery APIs, paid subscriptions, or field access. "
    "All analysis runs on free, publicly available data. The output is a ranked GeoJSON of "
    "candidate tiles and an interactive HTML discovery map, both generated automatically by "
    "the accompanying Python script (discovery_heatmap.py)."
)
body(
    "\nKey findings from the initial analysis:",
)
bullet("320 candidate tiles scored across four study regions")
bullet("Top candidates cluster in the Palpa–Nazca inter-valley zone (-14.55° to -14.58°S)")
bullet("All Tier 1 candidates (score ≥ 0.55) are within 0.2–2.7 km of a major river corridor")
bullet("The Ica–Pisco coastal corridor and Osmore Valley are the most under-surveyed high-potential zones")
bullet("North coast (Lambayeque–La Libertad) flagged as secondary priority pending additional site data")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# DATA FOUNDATION
# ═══════════════════════════════════════════════════════════════════════════
h1("1. Data Foundation")
h2("1.1 Known Site Inventory")
body(
    "The analysis is seeded by 34 verified archaeological sites extracted from 118 curated "
    "academic papers. Coordinates are manually verified from Wikipedia, the XRONOS database, "
    "and primary academic sources. Each site carries a confidence rating (high / medium) "
    "and coordinate status (exact / representative / approximate)."
)

# sites table
tbl = doc.add_table(rows=1, cols=4)
tbl.style = "Table Grid"
hdr = tbl.rows[0]
for i, txt in enumerate(["Site Name", "Lat/Lon", "Studies", "Confidence"]):
    hdr.cells[i].text = txt
    hdr.cells[i].paragraphs[0].runs[0].bold = True
shade_row(hdr, "1A237E")
for cell in hdr.cells:
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

site_data = [
    ("Nazca Lines",       "-14.6975, -75.1350", "61", "High"),
    ("Palpa Lines",       "-14.5966, -75.1949", "18", "High"),
    ("Ica",               "-14.0667, -75.7333", "9",  "High"),
    ("Cahuachi",          "-14.8186, -75.1167", "7",  "High"),
    ("Pinchango Alto",    "-14.4808, -75.1758", "5",  "Medium"),
    ("Acari Valley",      "-15.4311, -74.6158", "3",  "Medium"),
    ("Moquegua",          "-17.2000, -70.9333", "3",  "High"),
    ("Paracas Peninsula", "-13.8589, -76.3289", "3",  "High"),
    ("Rio Grande de Nazca","-14.5201, -75.2015","3",  "Medium"),
    ("Cerro Baul",        "-17.1121, -70.8588", "1",  "High"),
    ("Jauranga",          "-14.5466, -75.2047", "1",  "High"),
    ("Pernil Alto",       "-14.4920, -75.2160", "1",  "High"),
    ("Pampa Grande",      "-6.7626,  -79.4740", "1",  "High"),
    ("+ 21 additional sites", "…", "1 each", "High/Med"),
]
for i, row_data in enumerate(site_data):
    r = table_row(tbl, list(row_data))
    if i % 2 == 0:
        shade_row(r, "E8EAF6")
doc.add_paragraph()

h2("1.2 Academic Literature Coverage")
body(
    "Ten distinct search queries were executed across OpenAlex and Crossref, yielding "
    "373 raw papers. After deduplication and keyword curation, 118 papers remain covering "
    "remote sensing, geoglyphs, looting detection, settlement surveys, radiocarbon chronology, "
    "and LiDAR applications."
)
bullet("Most-studied site: Nazca Lines (61 papers)")
bullet("Least-studied high-potential zone: Ica–Pisco corridor, upper Nazca tributaries")
bullet("Notable recent paper: Snyder & Haas (2024) — satellite survey detects up to 20% of sites in coastal valleys")
bullet("Key method papers: Lasaponara & Masini (2018) space-based looting detection; Reindel & Gruen (2003) 3D Palpa modelling")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# STRATEGY LAYERS
# ═══════════════════════════════════════════════════════════════════════════
h1("2. Discovery Strategy: Six Predictive Layers")
body(
    "No single method reliably locates undiscovered sites. The strategy stacks six "
    "independent predictive layers into a composite score. Each layer is a probability "
    "surface over the study area. Cells where multiple layers converge — especially "
    "Layers 2 and 3 — represent the highest-priority investigation targets."
)

# weight table
tbl2 = doc.add_table(rows=1, cols=4)
tbl2.style = "Table Grid"
hdr2 = tbl2.rows[0]
for i, txt in enumerate(["Layer", "Name", "Weight", "Rationale"]):
    hdr2.cells[i].text = txt
    hdr2.cells[i].paragraphs[0].runs[0].bold = True
shade_row(hdr2, "1A237E")
for cell in hdr2.cells:
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

layer_weights = [
    ("1", "KDE Gap Analysis",          "15%", "Strongest near clusters; biased toward known areas"),
    ("2", "Hydro-Geomorphic Model",     "25%", "Strongest first-principles predictor — all sites need water"),
    ("3", "Spectral Anomaly Detection", "25%", "Direct physical evidence; best in arid desert"),
    ("4", "AI Vision / Geometry",       "15%", "High precision on structure detection; noisy at scale"),
    ("5", "Inter-Site Network",         "10%", "Theoretical; needs validation against ground truth"),
    ("6", "Research Gap Analysis",      "10%", "Meta-signal; identifies understudied areas"),
]
for i, row_data in enumerate(layer_weights):
    r = table_row(tbl2, list(row_data))
    if i % 2 == 0:
        shade_row(r, "E8EAF6")
doc.add_paragraph()

# ── Layer 1 ───────────────────────────────────────────────────────────────
h2("Layer 1: KDE Density + Gap Analysis")
body(
    "Apply Kernel Density Estimation to the 34 verified site coordinates using a bandwidth "
    "of ~100 km (approximately matching the spacing between major site clusters). This "
    "produces a continuous density surface over Peru's south coast. "
    "The 'gap score' targets cells that are within 2–15 km of a known site — "
    "close enough to be in an established archaeological zone, but far enough to be "
    "genuinely undiscovered territory."
)
bullet("Tool: scipy.stats.gaussian_kde, bandwidth=0.08°")
bullet("Sweet spot: 2–15 km from nearest known site")
bullet("Excludes: cells with a known site within 0.5 km")
bullet("Output: density-weighted gap probability per 1 km² cell")

h2("Layer 2: Hydro-Geomorphic Model (Highest Weight)")
body(
    "Ancient Peruvian civilizations — Nazca, Paracas, Wari, Tiwanaku — "
    "were hydraulic societies. Every major site clusters near a river or seasonal watercourse. "
    "This layer scores cells by proximity to the eleven encoded river corridors of Peru's "
    "south coast, including the Rio Grande de Nazca system and its five major tributaries."
)
bullet("Ideal proximity: 0.5–8 km from river centerline")
bullet("Scoring: linear decay 0–8 km, exponential decay beyond")
bullet("Rivers encoded: Rio Grande de Nazca, Palpa, Ingenio, Ica, Acari, Majes, Tambo, Osmore, Ilo, Las Trancas, Viscas")
bullet("Key insight: Upper tributary zones (headwaters of Ingenio, Aja, Las Trancas) are densely settled downstream but completely unsurveyed upstream")
bullet("Paleo-channels: Former riverbeds visible in SRTM DEM — adding these would upgrade this layer significantly")

h2("Layer 3: Satellite Spectral Anomaly Detection")
body(
    "Human occupation alters soil chemistry, moisture retention, and surface reflectance. "
    "Even fully buried structures are detectable in multispectral imagery, particularly in "
    "Peru's hyper-arid Nazca desert where vegetation is absent and contrast is extreme. "
    "This layer requires free Sentinel-2 imagery (10m resolution, 13 bands) — no API key "
    "required for manual download via Copernicus Open Access Hub."
)

# spectral index table
tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = "Table Grid"
hdr3 = tbl3.rows[0]
for i, txt in enumerate(["Index", "Formula", "What It Detects"]):
    hdr3.cells[i].text = txt
    hdr3.cells[i].paragraphs[0].runs[0].bold = True
shade_row(hdr3, "283EA7")
for cell in hdr3.cells:
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

spectral = [
    ("NDVI",          "(NIR−Red)/(NIR+Red)",               "Crop/vegetation marks over buried walls"),
    ("BSI",           "((SWIR+Red)−(NIR+Blue))/…",         "Exposed archaeological soils"),
    ("NDBI",          "(SWIR−NIR)/(SWIR+NIR)",             "Stone/adobe structural remnants"),
    ("Iron Oxide",    "Red / Blue",                         "Fired clay, ceramics, hearth remains"),
    ("Moisture Index","(NIR−SWIR)/(NIR+SWIR)",             "Subsurface walls affecting drainage"),
]
for i, row_data in enumerate(spectral):
    r = table_row(tbl3, list(row_data))
    if i % 2 == 0:
        shade_row(r, "E3F2FD")
doc.add_paragraph()

bullet("Best imagery: Sentinel-2 dry season composite (May–October), <10% cloud cover")
bullet("Method: compute all 5 indices, flag cells where 2+ indices deviate significantly from local neighborhood")
bullet("Peru advantage: hyper-arid conditions maximize spectral contrast of buried structures")
bullet("Validation: run on Cahuachi and Palpa Lines first — true positive rate confirms model calibration")

h2("Layer 4: AI Vision / Geometric Pattern Detection")
body(
    "Human construction produces geometric regularity that natural terrain lacks — "
    "linear alignments, rectangular enclosures, circular mounds, road traces, terrace "
    "edges, and irrigation channels. A multimodal vision model (Claude, GPT-4V) scanning "
    "satellite preview images can detect these patterns at scale."
)
bullet("Input: Google Earth or Sentinel-2 preview image per candidate tile")
bullet("Model: Claude claude-sonnet-4-6 or GPT-4V with structured JSON output")
bullet("Prompt: already built in prompts/vision_prompt.md")
bullet("Output: feature detections with type, confidence, approximate location")
bullet("Enhancement: provide 10–15 labelled examples of Nazca-area sites vs. natural terrain")
bullet("Human review: top 20 flagged tiles only — vision model narrows, humans decide")

h2("Layer 5: Inter-Site Network Analysis")
body(
    "Settlements form networks — trade routes, pilgrimage paths, administrative hierarchies. "
    "This layer models the expected connective structure of known sites and identifies "
    "locations where the network predicts a node but no site is recorded."
)
bullet("Method: Delaunay triangulation of known sites → least-cost path network (DEM-based)")
bullet("Flag: path intersections and midpoints on long routes (>30 km) with no site within 2 km")
bullet("Viewshed: hilltops visible from 3+ known sites → likely observation or signal posts")
bullet("Peru-specific gap: east–west connections between coastal and highland sites are under-represented in the data")

h2("Layer 6: Research Gap Analysis")
body(
    "Our 118 academic papers are not evenly distributed. Some zones (Nazca Lines: 61 papers) "
    "are extensively studied; others (Ica–Pisco corridor, upper tributaries) have near-zero "
    "coverage. Areas that are archaeologically plausible but academically ignored are "
    "disproportionately likely to contain undiscovered sites."
)
bullet("Method: geocode study areas from Peru_academic_master.csv, build 'research intensity' heatmap")
bullet("Priority: HIGH site-density zone + LOW research intensity = strong candidate")
bullet("Current gaps: inter-valley plateaus between Palpa and Nazca, upper Ingenio headwaters, Ica River upper reaches")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# COMPOSITE SCORING
# ═══════════════════════════════════════════════════════════════════════════
h1("3. Composite Scoring Formula")
body(
    "Each 1 km² cell receives a composite score between 0 and 1:"
)
formula_p = doc.add_paragraph()
formula_p.paragraph_format.left_indent = Inches(0.5)
run = formula_p.add_run(
    "DISCOVERY_SCORE = 0.40 × gap_score\n"
    "                + 0.30 × river_score\n"
    "                + 0.20 × coastal_zone_score\n"
    "                + 0.10 × density_context"
)
run.font.name = "Courier New"
run.font.size = Pt(10)

body(
    "\nCells with a known site within 0.5 km are zeroed out (already discovered). "
    "The weights above reflect the current no-satellite-API implementation. "
    "When Sentinel-2 spectral analysis is added, the recommended re-weighting is:"
)
bullet("Gap score: 0.15  (reduced — spectral data provides stronger signal)")
bullet("River score: 0.25  (unchanged)")
bullet("Spectral anomaly: 0.25  (new — direct physical evidence)")
bullet("Vision detections: 0.15  (new)")
bullet("Network node: 0.10  (new)")
bullet("Research gap: 0.10  (unchanged)")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# PRIORITY ZONES
# ═══════════════════════════════════════════════════════════════════════════
h1("4. Priority Search Zones")

h2("Zone A: Palpa–Nazca Inter-Valley Plateau (Highest Priority)")
body("Coordinates: -14.52° to -14.60°S, -75.17° to -75.21°W")
bullet("Why: Dense known sites to north (Palpa cluster) and south (Cahuachi, Nazca), plateau between them has no recorded sites")
bullet("Academic support: 79 papers on surrounding sites but near-zero coverage of the inter-valley zone")
bullet("River: Rio Grande system passes within 0.2–2.7 km of top-scored tiles")
bullet("Current score: 0.88 (top-ranked globally in our analysis)")
bullet("Recommended action: priority for Sentinel-2 spectral analysis and AI vision pass")

h2("Zone B: Upper Nazca Tributary Headwaters")
body("Coordinates: -14.30° to -14.70°S, -74.80° to -75.00°W")
bullet("Why: Main valleys are extensively studied; upper tributaries (Ingenio, Aja, Las Trancas) have zero academic coverage")
bullet("Ancient irrigation extended upstream — settlements would have followed")
bullet("Approximate size: 40 × 20 km → ~800 candidate tiles")
bullet("Recommended action: run hydro model on SRTM DEM to trace tributary channels; prioritize tiles at alluvial fans and confluences")

h2("Zone C: Ica–Pisco Coastal Corridor")
body("Coordinates: -13.70° to -14.10°S, -75.40° to -76.50°W")
bullet("Why: Only 4 verified sites for the entire zone; Paracas culture originated here — far more sites expected")
bullet("Academic gap: Paracas-related searches returned 99 papers, but coordinates cluster at peninsula only")
bullet("Approximate size: 40 × 50 km → requires hydro/spectral filtering to narrow to ~200 tiles")
bullet("Recommended action: highest priority for satellite spectral analysis; large flat coastal zone ideal for BSI/Iron Oxide indices")

h2("Zone D: Osmore Valley Side Tributaries")
body("Coordinates: -17.00° to -17.25°S, -70.50° to -71.10°W")
bullet("Why: 5 tightly clustered Wari-era sites; side valleys and ridgelines likely contain unrecorded observation posts and secondary settlements")
bullet("Network analysis applies strongly here — Cerro Baul is a Wari administrative center, satellite posts expected")
bullet("Approximate size: 20 × 20 km → ~400 candidate tiles")
bullet("Recommended action: viewshed analysis from Cerro Baul and Cerro Mejia; flag hilltops in line of sight")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# IMPLEMENTATION PLAN
# ═══════════════════════════════════════════════════════════════════════════
h1("5. Implementation Plan")

tbl4 = doc.add_table(rows=1, cols=4)
tbl4.style = "Table Grid"
hdr4 = tbl4.rows[0]
for i, txt in enumerate(["Step", "Action", "Effort", "Dependencies"]):
    hdr4.cells[i].text = txt
    hdr4.cells[i].paragraphs[0].runs[0].bold = True
shade_row(hdr4, "1A237E")
for cell in hdr4.cells:
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

steps = [
    ("1", "Populate sites_wgs84.geojson from verified academic coordinates", "1 hr", "None — data ready"),
    ("2", "Run Layer 1+2 gap + hydro analysis (discovery_heatmap.py)", "30 min", "Python: scipy, folium"),
    ("3", "Download SRTM 30m DEM via OpenTopography (no API key)", "1 hr", "Internet access"),
    ("4", "Refine hydro model with actual DEM drainage channels", "2 hr", "Step 3"),
    ("5", "Download Sentinel-2 composite via Copernicus Hub (free)", "2 hr", "Free ESA account"),
    ("6", "Run spectral indices (NDVI, BSI, NDBI, Iron Oxide)", "2 hr", "Steps 3+5, rasterio"),
    ("7", "Run AI vision on top 50 tiles from composite scoring", "1 hr", "Claude/GPT-4V API"),
    ("8", "Final composite scoring + map update", "1 hr", "Steps 2–7"),
]
for i, row_data in enumerate(steps):
    r = table_row(tbl4, list(row_data))
    if i % 2 == 0:
        shade_row(r, "E8EAF6")

doc.add_paragraph()
body(
    "Steps 1 and 2 are already complete. discovery_heatmap.py is operational and has "
    "generated 320 scored candidate tiles and an interactive HTML map. "
    "Steps 3–8 extend the analysis with terrain and satellite data when API access is available."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# OUTPUTS
# ═══════════════════════════════════════════════════════════════════════════
h1("6. Deliverables")
bullet("discovery_heatmap.html — interactive Folium map: known sites, KDE heatmap, river corridors, colour-coded candidate tiles (open in browser)")
bullet("candidate_tiles_scored.geojson — 320 scored 1 km² candidate tiles with composite score, gap/river/density sub-scores, and nearest known site/river distances")
bullet("Peru_academic_verified_coordinates.geojson — 34 verified site coordinates ready for GIS import")
bullet("Peru_academic_master.csv — 118 curated papers joined to site coordinates")
bullet("discovery_heatmap.py — fully reproducible analysis script (no API keys required)")

doc.add_paragraph()
h2("How to Use the Map")
bullet("Yellow dots = verified sites. Size proportional to academic study count.")
bullet("Heatmap layer = academic research density (red = heavily studied, blue = sparse)")
bullet("Blue polylines = river corridors. All high-scoring tiles should be near these.")
bullet("Red tiles (Tier 1, score ≥ 0.55) = highest priority for ground investigation")
bullet("Orange tiles (Tier 2, score 0.40–0.55) = secondary candidates")
bullet("Click any tile to see full score breakdown: gap, river, density scores + distances")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# CAVEATS
# ═══════════════════════════════════════════════════════════════════════════
h1("7. Limitations and Caveats")
bullet("River corridors are hand-encoded approximations. Upgrading to SRTM-derived drainage networks would significantly improve Layer 2 accuracy.")
bullet("Spectral anomaly layer (Layer 3) is not yet implemented — requires Sentinel-2 download. Without it, composite scores are based on gap + hydro only.")
bullet("34 known sites is a small training set. As more sites are added (from team contributions, YouTube reviews, community maps), KDE accuracy improves substantially.")
bullet("The North Coast (Lambayeque–La Libertad) scoring is weakest — only 1 site in that region (Pampa Grande) provides insufficient density context. Mochica-era sites are well-documented in literature but not yet in our dataset.")
bullet("Scores are relative, not absolute. A score of 0.88 means 'most promising in the current dataset' — not 'definitely contains a site'. Field verification is required.")
bullet("Coordinate status matters: 14 of 34 sites use 'representative' coordinates (district or valley level). As these are refined to exact coordinates, gap scoring will improve.")

doc.add_paragraph()
body(
    "\nPrepared by: Arkhub Team · NS Archaeology Hackathon · March 2026\n"
    "Data pipeline: OpenAlex, Crossref, XRONOS, Wikipedia · Python: scipy, folium, pandas",
)

doc.save(str(OUT))
print(f"Saved: {OUT}")
